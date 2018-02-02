# coding: utf-8
import os
import os.path
import fnmatch
import subprocess
import re
from collections import OrderedDict
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import cv2
import numpy as np

"""
直接调用shell
"""
#中英文写入规则不一样
languange='chi_sim_lstm'
show_box='false'
picture_format='jpg'
compen=0 #crop余量，有时候会超过图片范围

def image_None_rectangle_crop(image,width,point_list):
    mask = np.zeros(image.shape, dtype=np.uint8)
    coners_list=[(point[0],width-point[1]) for point in point_list]
    roi_corners = np.array([coners_list],dtype=np.int32)
    # fill the ROI so it doesn't get wiped out when the mask is applied
    if len(image.shape)==2:
        channel_count = 1  # i.e. 3 or 4 depending on your image
    else:
        channel_count=image.shape[-1]
    ignore_mask_color = (255,) * channel_count
    cv2.fillPoly(mask, roi_corners, ignore_mask_color)
    return cv2.bitwise_and(image, mask)

def image_crop(img,width,compensation,point_1,point_3):
    x1 = width - point_1[1] - compensation
    x2 = width - point_3[1] + compensation
    y1 = point_1[0] - compensation
    y2 = point_3[0] + compensation
    return img[x1:x2, y1:y2]


def filter_box(box_dict,height,width,filter_param_height,filter_param_width):
    """
    :param box_dict: raw box
    :param width:  image_width
    :param height:   image_height
    :return:
    """
    med_dic=box_dict.copy()
    print("raw boxes number: {}".format(len(med_dic)))
    new_dict={}
    for i in range(len(med_dic)):
        if len(med_dic[i])==5:
            point_1 = med_dic[i][0]
            point_3 = med_dic[i][2]
            type=med_dic[i][4]
        else:
            #后面可做连通图截取过程
            max_width=0
            min_width=10000
            max_height=0
            min_height=10000
            for point in med_dic[i][:-1]:
                if point[1]>max_width:
                    max_width=point[1]
                if point[1]<min_width:
                    min_width=point[1]
                if point[0]>max_height:
                    max_height=point[0]
                if point[0]<min_height:
                    min_height=point[0]
            point_1=(min_height,max_width)
            point_3=(max_height,min_width)
            type = med_dic[i][-1]

        box_width = point_1[1] - point_3[1]
        box_height = point_3[0] - point_1[0]
        #box只保留大段文字，图片，表格
        #针对表格和图片 width 垂直长度  height水平长度
        #这里只对规则的文本进行处理，多边形区域直接滤去了
        if type==19 or box_width<filter_param_width*width or box_height<filter_param_height*height:
            print("del box {}".format(i))
            del box_dict[i]
        else:
            #针对图片进行判断
            if type==3 and box_width<filter_param_width*5*width:
                print("del box {}".format(i))
                del box_dict[i]
        if i in box_dict.keys():
            new_dict[i] = [point_1, point_3, type]
    #对字典进行重排
    d = OrderedDict(sorted(new_dict.items(), key=lambda t: t[0]))
    d_raw=OrderedDict(sorted(box_dict.items(), key=lambda t: t[0]))
    return_dict={}
    return_dict_raw={}
    for i,item in enumerate(d):
    #重新编号字典
        return_dict[i]=d[item]
        return_dict_raw[i]=d_raw[item]
    print("final boxes number: {}".format(len(box_dict)))
    return return_dict,return_dict_raw

def crop_image(org_path,docx_path,result_path="./box_result.txt"):
    #write object
    document = Document()
    # load the box_result.txt and analysis
    box_list = []
    for line in open(result_path):
        box_list.append(line.replace('\n', '').split(' '))
    box_dict = {}
    for i, box in enumerate(box_list[:-1]):
        point_list = []
        for item in box[:-1]:
            item = item.split(',')
            point_list.append((int(item[0]), int(item[1])))
        point_list.append(int(box[-1]))
        box_dict[i] = point_list
    # 图片100的density识辨率更高？
    # box_dict  {num: [point point point point type]}
    # give the picture and ocr
    img = cv2.imread(org_path, -1)
    width, height = img.shape[0], img.shape[1]
    # fiter boxes
    #filter height and width
    #根据具体情况的字体来设置
    height_filter=0.1
    width_filter=0.005
    return_box,box_dict = filter_box(box_dict, height, width, height_filter,width_filter)
    for i in range(len(box_dict)):
        item=box_dict[i]
        if len(item)==5:
            point_1 = item[0]
            point_3 = item[2]
            type = item[4]
            crop = image_crop(img, width, compen, point_1, point_3)
        else:
            #求取多边形区域，目前没有找到切割算法，只能近似再切割，保证识别率最高
            crop=image_None_rectangle_crop(img,width,item[:-1])
            crop=image_crop(crop,width,compen,return_box[i][0],return_box[i][1])
            type=item[-1]
        crop_save_path="block{}.jpeg".format(i)

#picture preprocessing
        # crop = cv2.addWeighted(crop, 4, cv2.blur(crop, (30, 30)), -4, 128)
        # kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        # crop = cv2.filter2D(crop, -1, kernel)
        cv2.imwrite(crop_save_path, crop)
        #显示crop
        # cv2.imshow("image", crop)
        # cv2.waitKey(0)
        # 提取信息
        out_path=crop_save_path.replace('.jpeg','')
        #os.environ['SCROLLVIEW_PATH'] = "/Users/xiongfei/tesseract/java"
        if type==7 or type==6 or type==12 or type==11:
            subprocess.call(["tesseract", crop_save_path, out_path,
                                "-l",languange,
                             #分块识别，针对文字块，采用--psm 6最好
                                "--psm", "6",
                                "-c","Recognize_choice=True",
                                "-c", "textord_tabfind_show_blocks=false"])
            print("detection text, writing...")
            text_path=crop_save_path.replace('.jpeg','.txt')
            with open(text_path,'r') as f:
                #首先按照制表符切分段落
                all_the_text = f.read().strip().split("\n\n")
            #针对每一段的文字，对于特定的标题和文本内容，可以自行编辑正则和排列规则
            for item in all_the_text:
                if languange=='chi_sim_lstm':
                    item=''.join(item.split(' '))
                document.add_paragraph(item)
                #直接写入没问题，只是字体会有区别，设置标准单位Pt
                #document.paragraphs[-1].runs[0].font.size=Pt(10)
            print("complete writing {} paragraph".format(len(all_the_text)))
        elif type==8 or type==35:
            #对于表格先出去直线：主要要看脚本的参数形式
            #-negate replace every pixel with its complementary color
            #-morphology morphology analysis
            #-define
            subprocess.call(["convert",crop_save_path,
                             "-type","Grayscale",
                             "-negate",
                             "-define", "morphology:compose=darken",
                             "-morphology","Thinning",'Rectangle:1x80+0+0<',
                             "-negate",
                             crop_save_path])
            #纯文字识别 --psm 最好的效果，并且保持了文字的顺序
            os.environ['SCROLLVIEW_PATH'] = "/Users/xiongfei/tesseract/java"
            subprocess.call(["tesseract", crop_save_path, out_path,
                                "-l",languange,
                             #分块识别，针对文字块，采用--psm 6最好
                                "--psm", "6",
                                "-c","Recognize_choice=True",
                                "-c", "textord_tabfind_show_blocks=false"
                                # "-c", "plot_choice=True",
                                # "-c", "plot_choice_double=True",
                                # "-c", "plot_third=True"
                                ])
            print("detection table, writing...")
            table_path = crop_save_path.replace('.jpeg', '.txt')
            with open(table_path,'r') as f:
                all=f.read().strip().split('\n')
            #filter none line
            while '' in all:
                all.remove('')
            lens=len(all)
            table = document.add_table(lens, 1)
            for i in range(lens):
                table.cell(i,0).text=all[i]
            print("complete writing table")
            table.style='Table Grid'
        elif type==3 or type==4:
            print("detection image")
            document.add_picture(crop_save_path,width=Inches(1.25))
            print("complete writing Image")
    document.save(docx_path)
def execute(root_path):
    #os.walk存在隐藏文件，需要忽略
    for dirpath, _, filenames in os.walk(root_path):
        med=filenames.copy()
        for filename in med:
            if filename.startswith('.') or filename.endswith('.txt'):
                filenames.remove(filename)
        print(filenames)
        filenames.sort(key=lambda x: int(re.findall('(\d*).{}'.format(picture_format),x)[0]))
        for filename in filenames:
            if os.path.exists("./box_result.txt"):
                os.remove("./box_result.txt")
            print(filename)
            if fnmatch.fnmatch(filename, "*.{}".format(picture_format)):
                org_path = os.path.join(dirpath, filename)
                print("layout analysis {}".format(org_path))
                out_path=org_path.replace('.{}'.format(picture_format),'')
                os.environ['SCROLLVIEW_PATH'] = "/Users/xiongfei/tesseract/java"
                #--psm 3 检测表格的能力更强  --psm 4对单列文本分析更强，英文'eng'+--psm 4效果最好
                subprocess.call(["tesseract", org_path, out_path,
                                    "-l",languange,
                                    "-c","Recognize_choice=false",
                                    "-c", "textord_tabfind_show_blocks=True",
                                    "--psm","4",
                                    "-c", "plot_choice={}".format(show_box),
                                    "-c", "plot_choice_double={}.".format(show_box),
                                    "-c", "plot_third={}".format(show_box)])
                crop_image(org_path,filename.replace(picture_format,'docx'))
if __name__ == '__main__':
    #root_path = input("target folder path>")
    root_path="./convert_image/test"
    execute(root_path)